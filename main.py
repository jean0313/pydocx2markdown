from docx import Document as DocxDocument
import os

class WordExtractor:
    def __init__(self, file_path: str):
        self.file_path = file_path

    def _table_to_markdown(self, table, image_map):
        markdown = []
        # calculate the total number of columns
        total_cols = max(len(row.cells) for row in table.rows)

        header_row = table.rows[0]
        headers = self._parse_row(header_row, image_map, total_cols)
        markdown.append("| " + " | ".join(headers) + " |")
        markdown.append("| " + " | ".join(["---"] * total_cols) + " |")

        for row in table.rows[1:]:
            row_cells = self._parse_row(row, image_map, total_cols)
            markdown.append("| " + " | ".join(row_cells) + " |")
        return "\n".join(markdown)

    def _parse_row(self, row, image_map, total_cols):
        # Initialize a row, all of which are empty by default
        row_cells = [""] * total_cols
        col_index = 0
        for cell in row.cells:
            # make sure the col_index is not out of range
            while col_index < total_cols and row_cells[col_index] != "":
                col_index += 1
            # if col_index is out of range the loop is jumped
            if col_index >= total_cols:
                break
            cell_content = self._parse_cell(cell, image_map).strip()
            cell_colspan = cell.grid_span or 1
            for i in range(cell_colspan):
                if col_index + i < total_cols:
                    row_cells[col_index + i] = cell_content if i == 0 else ""
            col_index += cell_colspan
        return row_cells

    def _parse_cell(self, cell, image_map):
        cell_content = []
        for paragraph in cell.paragraphs:
            parsed_paragraph = self._parse_cell_paragraph(paragraph, image_map)
            if parsed_paragraph:
                cell_content.append(parsed_paragraph)
        unique_content = list(dict.fromkeys(cell_content))
        return " ".join(unique_content)
    
    def _parse_cell_paragraph(self, paragraph, image_map):
        paragraph_content = []
        for run in paragraph.runs:
            if run.element.xpath(".//a:blip"):
                for blip in run.element.xpath(".//a:blip"):
                    image_id = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                    if not image_id:
                        continue
                    image_part = paragraph.part.rels[image_id].target_part

                    if image_part in image_map:
                        image_link = image_map[image_part]
                        paragraph_content.append(image_link)
            else:
                paragraph_content.append(run.text)
        return "".join(paragraph_content).strip()
    
    def _parse_paragraph(self, paragraph, image_map):
        paragraph_content = []
        for run in paragraph.runs:
            if run.element.xpath(".//a:blip"):
                for blip in run.element.xpath(".//a:blip"):
                    embed_id = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                    if embed_id:
                        rel_target = run.part.rels[embed_id].target_ref
                        if rel_target in image_map:
                            paragraph_content.append(image_map[rel_target])
            if run.text.strip():
                paragraph_content.append(run.text.strip())
        return " ".join(paragraph_content) if paragraph_content else ""
    
    def parse_docx(self, docx_path, image_folder):
        doc = DocxDocument(docx_path)
        os.makedirs(image_folder, exist_ok=True)

        content = []
        image_map = {}

        def parse_paragraph(paragraph):
            paragraph_content = []
            for run in paragraph.runs:
                if hasattr(run.element, "tag") and isinstance(element.tag, str) and run.element.tag.endswith("r"):
                    drawing_elements = run.element.findall(
                        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
                    )
                    for drawing in drawing_elements:
                        blip_elements = drawing.findall(
                            ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
                        )
                        for blip in blip_elements:
                            embed_id = blip.get(
                                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                            )
                            if embed_id:
                                image_part = doc.part.related_parts.get(embed_id)
                                if image_part in image_map:
                                    paragraph_content.append(image_map[image_part])
                if run.text.strip():
                    paragraph_content.append(run.text.strip())
            return "".join(paragraph_content) if paragraph_content else ""

        paragraphs = doc.paragraphs.copy()
        tables = doc.tables.copy()
        for element in doc.element.body:
            if hasattr(element, "tag"):
                if isinstance(element.tag, str) and element.tag.endswith("p"):  # paragraph
                    para = paragraphs.pop(0)
                    parsed_paragraph = parse_paragraph(para)
                    if parsed_paragraph:
                        content.append(parsed_paragraph)
                elif isinstance(element.tag, str) and element.tag.endswith("tbl"):  # table
                    table = tables.pop(0)
                    content.append(self._table_to_markdown(table, image_map))
        return "\n".join(content)
    
    def extract(self):
        """Load given path as single page."""
        content = self.parse_docx(self.file_path, "storage")
        return content
    
extractor = WordExtractor("1.docx")
content = extractor.extract()

print(content)